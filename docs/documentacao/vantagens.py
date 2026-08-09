from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_heading("Piloto vs. Produção — Análise Técnica e Justificação da Migração", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Este documento explica por que razão o projeto evoluiu de uma versão Piloto, "
    "baseada em HTML/CSS/JavaScript e localStorage, para uma arquitetura de Produção "
    "baseada em Next.js, React, TypeScript, FastAPI e PostgreSQL. A análise considera "
    "segurança, dados, multiempresa, escalabilidade, manutenção, testes, IA, custos, "
    "vantagens, desvantagens e o balanço final da decisão."
)

doc.add_heading("1. Comparação das arquiteturas", level=1)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Camada"
hdr[1].text = "Piloto"
hdr[2].text = "Produção"

rows = [
    ("Interface", "HTML, CSS e JavaScript puro — 61 ficheiros HTML", "Next.js 16.3, React 19.2, TypeScript 5, Tailwind 4"),
    ("Lógica", "18 ficheiros JS em assets/js/", "React/TypeScript + FastAPI/Python"),
    ("Dados", "localStorage do browser", "PostgreSQL 18, 39 tabelas, 18 migrações"),
    ("Servidor", "Nenhum; ficheiros estáticos", "Backend FastAPI e infraestrutura de servidor"),
    ("Build", "Nenhum", "Build e execução de frontend/backend"),
    ("Dependências", "Nenhuma relevante", "Radix UI, Framer Motion, Recharts, SWR, big.js, SQLAlchemy, Pydantic, etc."),
    ("Autenticação", "Limitada pela natureza local", "PyJWT, bcrypt, pyotp, cryptography/Fernet"),
    ("Limites", "Não aplicável como serviço central", "SlowAPI"),
    ("IA", "Não fazia parte da arquitetura original", "OpenAI, exclusivamente para o assistente"),
    ("Testes", "Limitados", "pytest — 359 testes"),
    ("Qualidade", "Sem pipeline moderno de tipagem", "Biome + TypeScript strict"),
]
for r in rows:
    cells = table.add_row().cells
    for i, value in enumerate(r):
        cells[i].text = value

doc.add_heading("2. Por que foi necessário sair do Piloto?", level=1)
doc.add_paragraph(
    "O Piloto não foi uma má escolha. Pelo contrário, foi adequado para validar rapidamente "
    "a ideia, testar fluxos e construir uma primeira versão funcional com baixo custo e pouca "
    "infraestrutura. A combinação de HTML, CSS, JavaScript puro e localStorage permitiu executar "
    "a aplicação sem servidor, base de dados ou processo de deployment."
)
doc.add_paragraph(
    "O problema surge quando o objetivo deixa de ser apenas demonstrar ou validar o conceito "
    "e passa a ser operar uma plataforma empresarial real. O localStorage pertence ao navegador "
    "de cada utilizador, não oferece uma base centralizada de dados e não é uma solução adequada "
    "para uma aplicação multiutilizador e multiempresa com requisitos de segurança, auditoria, "
    "backup, concorrência e persistência."
)
doc.add_paragraph(
    "Assim, a mudança para Produção não aconteceu porque o Piloto estava errado, mas porque "
    "o objetivo do sistema mudou: passou de uma prova de conceito funcional para uma plataforma "
    "preparada para utilização real."
)

doc.add_heading("3. Segurança", level=1)
doc.add_paragraph(
    "No Piloto, uma parte importante da lógica e dos dados está no próprio navegador. Isso limita "
    "a capacidade de controlar centralmente quem pode executar cada operação. Na Produção, o backend "
    "pode aplicar autenticação, autorização, perfis, permissões, validação dos dados, regras de negócio, "
    "limites e auditoria no servidor."
)
doc.add_paragraph(
    "A Produção utiliza PyJWT, bcrypt, pyotp e cryptography/Fernet, permitindo construir mecanismos "
    "mais robustos de autenticação, proteção de palavras-passe, segundo fator e proteção criptográfica "
    "de dados. Isto não significa que a Produção seja automaticamente segura: a configuração correta, "
    "gestão de segredos, HTTPS, firewall, backups, atualizações e políticas de acesso continuam sendo "
    "responsabilidades essenciais."
)

doc.add_heading("4. Base de dados e integridade dos dados", level=1)
doc.add_paragraph(
    "A principal mudança é a substituição do armazenamento local do navegador por PostgreSQL 18, "
    "com 39 tabelas e 18 migrações. Uma base de dados relacional permite centralização, relações, "
    "constraints, índices, transações, integridade referencial, concorrência, backups, recuperação "
    "e evolução controlada através de migrations."
)
doc.add_paragraph(
    "Para um sistema contabilístico, esta mudança é particularmente importante porque os dados "
    "financeiros não devem depender do armazenamento local de um único navegador."
)

doc.add_heading("5. Multiempresa e multiutilizador", level=1)
doc.add_paragraph(
    "O Piloto é naturalmente próximo de uma aplicação executada localmente no navegador. "
    "A Produção pode operar como uma plataforma centralizada onde várias empresas e utilizadores "
    "trabalham simultaneamente, com isolamento e controlo de acesso por empresa."
)
doc.add_paragraph(
    "A arquitetura permite organizar a plataforma em entidades como empresas, utilizadores, "
    "exercícios, diários, lançamentos e uma área administrativa, mantendo os dados no servidor."
)

doc.add_heading("6. Escalabilidade", level=1)
doc.add_paragraph(
    "Enquanto o Piloto é adequado para uma utilização pequena e local, a Produção foi desenhada "
    "para crescer com maior quantidade de utilizadores, empresas, lançamentos, histórico, relatórios "
    "e operações simultâneas. PostgreSQL e um backend dedicado oferecem uma base muito mais apropriada "
    "para esse crescimento."
)

doc.add_heading("7. Manutenção e evolução", level=1)
doc.add_paragraph(
    "O Piloto possui 61 ficheiros HTML e 18 ficheiros JavaScript. Essa abordagem é simples para começar, "
    "mas torna-se mais difícil de manter à medida que a aplicação cresce. A Produção utiliza React, "
    "TypeScript, componentes reutilizáveis, Tailwind, SWR e outras ferramentas modernas, permitindo "
    "uma organização mais estruturada e deteção antecipada de determinados erros."
)

doc.add_heading("8. Backend e regras de negócio", level=1)
doc.add_paragraph(
    "No Piloto, grande parte da lógica executa no navegador. Na Produção, o fluxo passa pelo backend: "
    "o frontend comunica com a API, o backend autentica o utilizador, aplica autorização, valida os dados, "
    "executa as regras de negócio e persiste os resultados na base de dados."
)
doc.add_paragraph(
    "Isso é particularmente importante para regras contabilísticas, porque uma regra crítica deve ser "
    "aplicada no servidor e não depender apenas da interface apresentada ao utilizador."
)

doc.add_heading("9. Auditoria", level=1)
doc.add_paragraph(
    "A arquitetura de Produção permite criar uma camada de auditoria muito mais completa, registando "
    "operações importantes com informação como utilizador, empresa, operação e momento da ação. "
    "Isso é essencial para uma plataforma empresarial em que as alterações precisam deixar rastreio."
)

doc.add_heading("10. IA e controlo de custos", level=1)
doc.add_paragraph(
    "A Produção introduz integração com a OpenAI exclusivamente para o assistente. Como essa integração "
    "fica no backend, a plataforma pode centralizar a chave de API, controlar o modelo, limitar tokens, "
    "registar consumo e calcular custos por empresa."
)
doc.add_paragraph(
    "A arquitetura também permite que o superadmin controle o modelo utilizado, o limite máximo de tokens "
    "de saída e os preços usados para estimativas de custo. Os preços devem ser configuráveis e não ficar "
    "fixos permanentemente no código, pois podem mudar."
)

doc.add_heading("11. Testes e qualidade", level=1)
doc.add_paragraph(
    "A Produção possui 359 testes com pytest, além de TypeScript strict e Biome. Isso não significa que "
    "o sistema seja perfeito, mas fornece uma base automatizada para detetar regressões e manter o comportamento "
    "esperado durante a evolução."
)

doc.add_heading("12. Principais vantagens da Produção", level=1)
advantages = [
    "Dados centralizados e persistentes.",
    "Base de dados relacional PostgreSQL.",
    "Autenticação e autorização no backend.",
    "Possibilidade de gestão multiempresa e multiutilizador.",
    "Maior capacidade de auditoria.",
    "Maior capacidade de escalabilidade.",
    "Backend dedicado para aplicar regras de negócio.",
    "Melhor organização e manutenção do frontend.",
    "Tipagem com TypeScript strict.",
    "Testes automatizados.",
    "Migrations controladas.",
    "Integração segura da IA pelo backend.",
    "Controlo de tokens e custos.",
    "Preparação para deployment e operação em servidor."
]
for item in advantages:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("13. Desvantagens da Produção", level=1)
disadvantages = [
    "Maior complexidade técnica.",
    "Necessidade de servidor e infraestrutura.",
    "Custos de hospedagem e manutenção.",
    "Necessidade de backups e monitorização.",
    "Maior superfície de ataque.",
    "Necessidade de configuração correta de HTTPS, firewall, CORS e segredos.",
    "Mais componentes que precisam ser mantidos compatíveis.",
    "Deploy e atualizações são mais complexos que simplesmente abrir ficheiros estáticos.",
    "Maior esforço inicial de desenvolvimento e manutenção."
]
for item in disadvantages:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("14. O Piloto continua a ter valor?", level=1)
doc.add_paragraph(
    "Sim. O Piloto não deve ser descartado. Ele continua sendo uma referência funcional importante "
    "para validar fluxos e regras que já foram testados. A Produção deve preservar as regras e funcionalidades "
    "validadas no Piloto, enquanto acrescenta a infraestrutura necessária para uma operação real."
)

doc.add_heading("15. Balanço final: compensou a migração?", level=1)
doc.add_paragraph(
    "Para uma demonstração ou prova de conceito, o Piloto seria suficiente e apresentaria vantagens claras "
    "de simplicidade, velocidade e baixo custo. Para uma plataforma empresarial real, porém, as limitações "
    "estruturais do localStorage e da ausência de backend e base de dados central tornam a arquitetura inadequada "
    "como solução definitiva."
)
doc.add_paragraph(
    "Por isso, considerando o objetivo final de disponibilizar uma plataforma real para várias empresas e "
    "utilizadores, com dados centralizados, autenticação, segurança, auditoria, escalabilidade, IA e capacidade "
    "de operar num servidor, a migração para Produção compensou. O custo foi o aumento da complexidade e da "
    "necessidade de infraestrutura, mas esse é o custo natural de transformar uma prova de conceito numa "
    "plataforma empresarial."
)

doc.add_heading("16. Síntese comparativa", level=1)
table2 = doc.add_table(rows=1, cols=3)
table2.style = "Table Grid"
h = table2.rows[0].cells
h[0].text = "Critério"
h[1].text = "Piloto"
h[2].text = "Produção"
comparison = [
    ("Simplicidade", "Muito alta", "Média"),
    ("Desenvolvimento inicial", "Muito rápido", "Mais complexo"),
    ("Execução local", "Excelente", "Boa, mas depende de serviços"),
    ("Segurança", "Limitada", "Muito mais adequada"),
    ("Dados centralizados", "Não", "Sim"),
    ("Multiutilizador", "Limitado", "Sim"),
    ("Multiempresa", "Limitado", "Preparado"),
    ("Base de dados real", "Não", "PostgreSQL"),
    ("Auditoria", "Limitada", "Muito mais adequada"),
    ("Escalabilidade", "Baixa", "Muito maior"),
    ("Manutenção profissional", "Limitada", "Muito maior"),
    ("Testes automatizados", "Limitados", "359 testes"),
    ("Tipagem", "JavaScript", "TypeScript strict"),
    ("Backend/API", "Não", "FastAPI"),
    ("IA", "Não na arquitetura original", "OpenAI no assistente"),
    ("Preparação para produto real", "Baixa/média", "Alta")
]
for r in comparison:
    cells = table2.add_row().cells
    for i, value in enumerate(r):
        cells[i].text = value

doc.add_heading("Conclusão", level=1)
doc.add_paragraph(
    "O Piloto foi a prova de conceito; a Produção representa a transformação dessa prova de conceito "
    "numa arquitetura preparada para operar como produto real. A migração não deve ser entendida como "
    "uma rejeição do Piloto, mas como uma evolução necessária perante novos requisitos de segurança, "
    "persistência, multiempresa, escalabilidade, auditoria, testes e operação em servidor."
)
doc.add_paragraph(
    "A principal vantagem da Produção é que ela cria uma base técnica sobre a qual a plataforma pode "
    "crescer de forma controlada. A principal desvantagem é a complexidade adicional. Para o objetivo "
    "empresarial definido, o balanço é claramente favorável à Produção."
)

path = "/mnt/data/Analise_Piloto_vs_Producao.docx"
doc.save(path)
path
